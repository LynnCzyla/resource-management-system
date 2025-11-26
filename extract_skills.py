import os
import re
import tempfile
import logging
from fastapi import APIRouter, UploadFile, File
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from typing import List
from io import BytesIO
from PIL import Image
import spacy

# ---------- CONFIG ----------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler-24.08.0\Library\bin"
nlp = spacy.load("en_core_web_sm")

# ---------- LOGGING CONFIG ----------
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger("ocr_nlp_logger")

router = APIRouter()

# ---------- SKILLS / FRAMEWORKS ----------
SKILL_KEYWORDS = [
    "Python", "Java", "JavaScript", "React", "SQL", "NoSQL",
    "API Design", "Project Management", "Agile", "Scrum", "System Architecture",
    "Leadership", "Teamwork", "Communication", "Machine Learning"
]

SKILL_SYNONYMS = {
    "JS": "JavaScript",
    "PM": "Project Management",
    "ML": "Machine Learning",
    "DB": "SQL",
    "PyTorch": "PyTorch",
}

COMMON_FRAMEWORKS = [
    "Django", "Flask", "FastAPI", "Vue", "Angular",
    "Node.js", "Express", "Spring Boot"
]

TECH_TERMS = set([
    "Python", "Java", "JavaScript", "React", "SQL", "NoSQL",
    "Django", "Flask", "FastAPI", "Vue", "Angular", "Node.js", "Express",
    "Spring Boot", "PyTorch", "TensorFlow", "Keras", "Docker", "Kubernetes",
    "Terraform", "AWS", "GCP", "Azure", "MongoDB", "PostgreSQL", "MySQL"
])


# ------------------------------------------------------
#   OCR: PDF TEXT EXTRACTION
# ------------------------------------------------------
def extract_text_from_pdf(pdf_path):
    logger.info(f"Starting PDF OCR extraction for: {pdf_path}")
    text = ""

    try:
        images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)
        logger.info(f"PDF converted into {len(images)} image pages")

        for i, img in enumerate(images):
            logger.debug(f"OCR reading page {i + 1}")
            page_text = pytesseract.image_to_string(img)
            logger.debug(f"OCR text length from page {i + 1}: {len(page_text)}")
            text += page_text + "\n"

    except Exception as e:
        logger.error(f"PDF OCR extraction failed: {e}")

    logger.info(f"Finished PDF OCR extraction. Total characters: {len(text)}")
    return text.strip()


# ------------------------------------------------------
#   DOCX TEXT EXTRACTION
# ------------------------------------------------------
def extract_text_from_docx(docx_file):
    logger.info("Starting DOCX text extraction")
    text = ""

    try:
        doc = Document(docx_file)

        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")

    logger.info(f"Finished DOCX extraction. Characters: {len(text)}")
    return text


# ------------------------------------------------------
#   TEXT CLEANING
# ------------------------------------------------------
def clean_ocr_text(text):
    logger.debug("Cleaning extracted text")
    headings = [
        "Personal Information", "Education", "Skills", "Experience",
        "Work Experience", "Projects", "Certifications", "Summary"
    ]

    for h in headings:
        text = text.replace(h, "")

    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


# ------------------------------------------------------
#   PERSONAL INFO EXTRACTION
# ------------------------------------------------------
def extract_personal_info(text):
    logger.info("Extracting personal info")
    info = {}
    cleaned = clean_ocr_text(text)
    lines = cleaned.split("\n")

    for line in lines:
        line_lower = line.lower()

        if "full name" in line_lower:
            match = re.search(r"Full Name[:\s]*(.+)", line)
            if match:
                info["Full Name"] = match.group(1).strip()

        elif "employee id" in line_lower or re.search(r"\bID\b", line):
            match = re.search(r"(Employee ID|ID)[:\s]*(.+)", line)
            if match:
                info["Employee ID"] = match.group(2).strip()

        elif "email" in line_lower:
            match = re.search(r"[\w\.-]+@[\w\.-]+", line)
            if match:
                info["Email"] = match.group()

        elif "phone" in line_lower or re.search(r"\+?\d{9,}", line):
            match = re.search(r"(\+?\d[\d\s\-]{8,}\d)", line)
            if match:
                info["Phone Number"] = match.group()

        elif "location" in line_lower:
            match = re.search(r"Location[:\s]*(.+)", line)
            if match:
                info["Location"] = match.group(1).strip()

    logger.debug("Applying NLP fallback for personal info")
    doc = nlp(cleaned)

    if "Full Name" not in info:
        names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if names:
            info["Full Name"] = names[0]

    if "Location" not in info:
        locations = [ent.text for ent in doc.ents if ent.label_ == "GPE"]
        if locations:
            info["Location"] = locations[0]

    if "Organization" not in info:
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        if orgs:
            info["Organization"] = orgs[0]

    logger.info(f"Extracted personal info: {info}")
    return info


# ------------------------------------------------------
#   SKILL EXTRACTION
# ------------------------------------------------------
def extract_skills_dynamic(text):
    logger.info("Starting skill extraction")
    found_skills = set()
    text_lower = text.lower()

    # Keyword skills
    for skill in SKILL_KEYWORDS:
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
            logger.debug(f"Skill found (keyword): {skill}")
            found_skills.add(skill)

    # Frameworks
    for fw in COMMON_FRAMEWORKS:
        if re.search(r'\b' + re.escape(fw.lower()) + r'\b', text_lower):
            logger.debug(f"Skill found (framework): {fw}")
            found_skills.add(fw)

    # Synonyms
    for syn, actual in SKILL_SYNONYMS.items():
        if re.search(r'\b' + re.escape(syn.lower()) + r'\b', text_lower):
            logger.debug(f"Skill found (synonym {syn} → {actual})")
            found_skills.add(actual)

    # NLP fallback
    doc = nlp(text)
    for token in doc:
        if token.text.strip() in TECH_TERMS:
            logger.debug(f"Skill found (NLP): {token.text.strip()}")
            found_skills.add(token.text.strip())

    result = sorted(found_skills)
    logger.info(f"Final extracted skills: {result}")
    return result


# ------------------------------------------------------
#   API ROUTE
# ------------------------------------------------------
@router.post("/extract_skills/")
async def extract_skills_endpoint(files: List[UploadFile] = File(...)):
    logger.info("API /extract_skills called")
    results = []

    for file in files:
        logger.info(f"Processing uploaded file: {file.filename}")
        content = await file.read()
        suffix = os.path.splitext(file.filename)[1].lower()
        text = ""

        try:
            if suffix == ".pdf":
                logger.info(f"Handling PDF file: {file.filename}")

                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    tmp_pdf.write(content)
                    tmp_pdf.flush()
                    pdf_path = tmp_pdf.name

                text = extract_text_from_pdf(pdf_path)
                os.unlink(pdf_path)

            elif suffix == ".docx":
                logger.info(f"Handling DOCX file: {file.filename}")
                docx_file = BytesIO(content)
                text = extract_text_from_docx(docx_file)

            elif suffix in [".png", ".jpg", ".jpeg"]:
                logger.info(f"Handling image file: {file.filename}")
                img = Image.open(BytesIO(content))
                text = pytesseract.image_to_string(img)
                logger.debug(f"Image OCR text length: {len(text)}")

            else:
                logger.warning(f"Unsupported file type: {suffix}")
                continue

            if not text.strip():
                logger.warning(f"No text extracted from file: {file.filename}")

            personal_info = extract_personal_info(text)
            skills = extract_skills_dynamic(text)

            results.append({
                "filename": file.filename,
                "personal_info": personal_info,
                "skills": skills
            })

        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {e}")

    all_skills = sorted(set(skill for r in results for skill in r["skills"]))

    logger.info(f"Returning response. Total unique skills: {all_skills}")
    return {"results": results, "skills": all_skills}
